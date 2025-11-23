"""
Document parser for extracting text from PDF, DOCX, and TXT files.
Supports intelligent chunking by sections and paragraphs.
"""

import os
from pathlib import Path
from typing import List, Dict, Optional
import PyPDF2
import pdfplumber
from docx import Document
import re


class DocumentChunk:
    """Represents a chunk of text with metadata."""
    
    def __init__(self, text: str, metadata: Dict[str, str]):
        self.text = text
        self.metadata = metadata
    
    def __repr__(self):
        return f"DocumentChunk(text='{self.text[:50]}...', metadata={self.metadata})"


class DocumentParser:
    """Parse documents and extract text with intelligent chunking."""
    
    def __init__(self, chunk_size: int = 500, chunk_overlap: int = 50):
        """
        Initialize parser with chunking parameters.
        
        Args:
            chunk_size: Target size for each chunk in characters
            chunk_overlap: Number of characters to overlap between chunks
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
    
    def parse_file(self, file_path: str) -> List[DocumentChunk]:
        """
        Parse a file and return chunks with metadata.
        
        Args:
            file_path: Path to the file to parse
            
        Returns:
            List of DocumentChunk objects
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Determine file type and parse accordingly
        ext = file_path.suffix.lower()
        
        if ext == '.pdf':
            text = self._parse_pdf(file_path)
        elif ext == '.docx':
            text = self._parse_docx(file_path)
        elif ext == '.txt':
            text = self._parse_txt(file_path)
        else:
            raise ValueError(f"Unsupported file type: {ext}")
        
        # Create chunks with metadata
        chunks = self._create_chunks(text, file_path)
        
        return chunks
    
    def _parse_pdf(self, file_path: Path) -> str:
        """Extract text from PDF file."""
        text = ""
        
        # Try pdfplumber first (better for complex PDFs)
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception as e:
            print(f"pdfplumber failed, trying PyPDF2: {e}")
            
            # Fallback to PyPDF2
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page in pdf_reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
            except Exception as e2:
                raise Exception(f"Failed to parse PDF: {e2}")
        
        return text.strip()
    
    def _parse_docx(self, file_path: Path) -> str:
        """Extract text from DOCX file."""
        try:
            doc = Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + " "
                text += "\n"
            
            return text.strip()
        except Exception as e:
            raise Exception(f"Failed to parse DOCX: {e}")
    
    def _parse_txt(self, file_path: Path) -> str:
        """Extract text from TXT file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read().strip()
        except UnicodeDecodeError:
            # Try different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    return file.read().strip()
            except Exception as e:
                raise Exception(f"Failed to parse TXT: {e}")
    
    def _create_chunks(self, text: str, file_path: Path) -> List[DocumentChunk]:
        """
        Create intelligent chunks from text.
        
        Tries to split on section boundaries, then paragraphs, then sentences.
        """
        chunks = []
        
        # First, try to identify sections (headers like "Experience", "Education", etc.)
        sections = self._split_into_sections(text)
        
        for section_name, section_text in sections:
            # If section is small enough, keep it as one chunk
            if len(section_text) <= self.chunk_size:
                chunks.append(DocumentChunk(
                    text=section_text,
                    metadata={
                        'filename': file_path.name,
                        'filepath': str(file_path),
                        'section': section_name,
                        'type': 'section'
                    }
                ))
            else:
                # Split large sections into smaller chunks
                section_chunks = self._split_text(section_text)
                for i, chunk_text in enumerate(section_chunks):
                    chunks.append(DocumentChunk(
                        text=chunk_text,
                        metadata={
                            'filename': file_path.name,
                            'filepath': str(file_path),
                            'section': section_name,
                            'chunk_index': i,
                            'type': 'chunk'
                        }
                    ))
        
        return chunks
    
    def _split_into_sections(self, text: str) -> List[tuple]:
        """
        Split text into sections based on common resume headers.
        
        Returns:
            List of (section_name, section_text) tuples
        """
        # Common resume section headers
        section_patterns = [
            r'^(SUMMARY|OBJECTIVE|PROFILE)',
            r'^(EXPERIENCE|WORK EXPERIENCE|PROFESSIONAL EXPERIENCE|EMPLOYMENT)',
            r'^(EDUCATION|ACADEMIC BACKGROUND)',
            r'^(SKILLS|TECHNICAL SKILLS|CORE COMPETENCIES)',
            r'^(PROJECTS|NOTABLE PROJECTS)',
            r'^(CERTIFICATIONS|LICENSES)',
            r'^(AWARDS|ACHIEVEMENTS|HONORS)',
            r'^(PUBLICATIONS|RESEARCH)',
            r'^(CONTACT|CONTACT INFORMATION)',
        ]
        
        lines = text.split('\n')
        sections = []
        current_section = 'Header'
        current_text = []
        
        for line in lines:
            line_upper = line.strip().upper()
            
            # Check if this line is a section header
            is_header = False
            for pattern in section_patterns:
                if re.match(pattern, line_upper):
                    # Save previous section
                    if current_text:
                        sections.append((current_section, '\n'.join(current_text).strip()))
                    
                    current_section = line.strip()
                    current_text = []
                    is_header = True
                    break
            
            if not is_header:
                current_text.append(line)
        
        # Add final section
        if current_text:
            sections.append((current_section, '\n'.join(current_text).strip()))
        
        return sections if sections else [('Document', text)]
    
    def _split_text(self, text: str) -> List[str]:
        """
        Split text into chunks with overlap.
        
        Tries to split on paragraph boundaries, then sentences.
        """
        chunks = []
        
        # Split into paragraphs
        paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
        
        if not paragraphs:
            paragraphs = [p.strip() for p in text.split('\n') if p.strip()]
        
        current_chunk = ""
        
        for para in paragraphs:
            # If adding this paragraph would exceed chunk size
            if len(current_chunk) + len(para) > self.chunk_size and current_chunk:
                chunks.append(current_chunk.strip())
                
                # Start new chunk with overlap
                overlap_text = current_chunk[-self.chunk_overlap:] if len(current_chunk) > self.chunk_overlap else current_chunk
                current_chunk = overlap_text + " " + para
            else:
                current_chunk += " " + para if current_chunk else para
        
        # Add final chunk
        if current_chunk:
            chunks.append(current_chunk.strip())
        
        return chunks if chunks else [text]


def parse_document(file_path: str, chunk_size: int = 500, chunk_overlap: int = 50) -> List[DocumentChunk]:
    """
    Convenience function to parse a document.
    
    Args:
        file_path: Path to the document
        chunk_size: Target size for each chunk
        chunk_overlap: Overlap between chunks
        
    Returns:
        List of DocumentChunk objects
    """
    parser = DocumentParser(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    return parser.parse_file(file_path)

